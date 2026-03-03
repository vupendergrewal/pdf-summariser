from flask import Flask, request, render_template, send_file
from werkzeug.utils import secure_filename
import fitz
import nltk
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
import os
import io
import json
from groq import Groq

nltk.download('punkt', quiet=True)
nltk.download('punkt_tab', quiet=True)

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs('uploads', exist_ok=True)

client = Groq(api_key=os.environ.get("GROQ_API_KEY"))

def extract_text(pdf_path):
    doc = fitz.open(pdf_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def summarise_text(text, num_sentences=7):
    parser = PlaintextParser.from_string(text, Tokenizer("english"))
    summarizer = LsaSummarizer()
    summary = summarizer(parser.document, num_sentences)
    return " ".join(str(s) for s in summary)

def get_reading_time(text):
    words = len(text.split())
    minutes = round(words / 200)
    return max(1, minutes)

def get_ai_insights(text):
    try:
        response = client.chat.completions.create(
          model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": """Analyze this document and return ONLY a JSON object with these exact keys:
- sentiment: one word (Positive, Negative, or Neutral)
- topics: list of 5 short key topics
- difficulty: one word (Easy, Moderate, or Complex)
No extra text, just the JSON."""},
                {"role": "user", "content": f"Document:\n{text[:4000]}"}
            ]
        )
        raw = response.choices[0].message.content
        clean = raw.replace("```json", "").replace("```", "").strip()
        return json.loads(clean)
    except:
        return {
            "sentiment": "Neutral",
            "topics": ["Topic 1", "Topic 2", "Topic 3", "Topic 4", "Topic 5"],
            "difficulty": "Moderate"
        }

@app.route("/")
def home():
    return render_template("index.html")

@app.route("/summarise", methods=["POST"])
def summarise():
    file = request.files['pdf']
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)

    text = extract_text(filepath)

    if not text.strip():
        return render_template("result.html",
            error="No text found. PDF might be scanned.",
            summary=None, filename=None,
            word_count=0, sentence_count=0,
            document_text=None, reading_time=0,
            sentiment="Neutral", topics=[], difficulty="Moderate")

    summary = summarise_text(text)
    word_count = len(text.split())
    sentence_count = summary.count('.') + 1
    reading_time = get_reading_time(text)
    insights = get_ai_insights(text)

    return render_template("result.html",
        summary=summary,
        filename=filename,
        word_count=word_count,
        sentence_count=sentence_count,
        document_text=text[:8000],
        reading_time=reading_time,
        sentiment=insights["sentiment"],
        topics=insights["topics"],
        difficulty=insights["difficulty"],
        error=None)

@app.route("/download/word", methods=["POST"])
def download_word():
    summary = request.form.get("summary")
    filename = request.form.get("filename")
    doc = Document()
    doc.add_heading("PDF Summary", 0)
    doc.add_paragraph(f"Original file: {filename}")
    doc.add_paragraph(" ")
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(summary)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True,
        download_name="summary.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.route("/download/pdf", methods=["POST"])
def download_pdf():
    summary = request.form.get("summary")
    filename = request.form.get("filename")
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    c.setFont("Helvetica-Bold", 20)
    c.drawString(50, height - 60, "PDF Summary")
    c.setFont("Helvetica", 11)
    c.drawString(50, height - 90, f"Original file: {filename}")
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, height - 130, "Summary:")
    c.setFont("Helvetica", 11)
    y = height - 160
    words = summary.split()
    line = ""
    for word in words:
        if c.stringWidth(line + word, "Helvetica", 11) < width - 100:
            line += word + " "
        else:
            c.drawString(50, y, line)
            y -= 20
            line = word + " "
            if y < 50:
                c.showPage()
                y = height - 50
    c.drawString(50, y, line)
    c.save()
    buffer.seek(0)
    return send_file(buffer, as_attachment=True,
        download_name="summary.pdf",
        mimetype="application/pdf")

@app.route("/chat", methods=["POST"])
def chat():
    data = request.get_json()
    question = data.get("question")
    document_text = data.get("document_text", "")
    history = data.get("history", [])

    if not document_text:
        return {"answer": "Document missing. Please upload again."}

    try:
        messages = [
            {"role": "system", "content": f"You are a helpful assistant. Answer questions based only on this document:\n\n{document_text[:6000]}"}
        ]

        for h in history[-6:]:
            messages.append({"role": "user", "content": h["question"]})
            messages.append({"role": "assistant", "content": h["answer"]})

        messages.append({"role": "user", "content": question})

        response = client.chat.completions.create(
           model="llama-3.3-70b-versatile",
            messages=messages
        )
        return {"answer": response.choices[0].message.content}
    except Exception as e:
        return {"answer": f"Error: {str(e)}"}

if __name__ == "__main__":
    app.run(debug=True)